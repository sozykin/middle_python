from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Документ, созданный программой Python', level=1)

p = document.add_paragraph()

p = document.add_paragraph('Текст с выделением ')
p.add_run('жирным шрифтом').bold = True
p.add_run(' и ')
p.add_run('наклонным шрифтом.').italic = True

p1 = document.add_paragraph()
p1.add_run('Сетевые протоколы: ').bold = True

document.add_paragraph(
    'HTTP - протокол передачи гипертекста.', style='List Bullet'
)
document.add_paragraph(
    'STMP - простой протокол передачи почты.', style='List Bullet'
)

document.save('demo.docx')